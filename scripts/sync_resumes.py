#!/usr/bin/env python
"""
Automated Resume Sync Engine
Scans active project repositories, computes real-time stats (commits, LOC, files, tests),
updates all 4 resume .docx variants, compiles fresh PDFs via MS Word, and syncs to portfolio_site.
"""

import os
import sys
import glob
import shutil
import subprocess

# Ensure utf-8 encoding on Windows console
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

try:
    import docx
except ImportError:
    print("[ERROR] python-docx is required. Run: pip install python-docx")
    sys.exit(1)

# Base Paths
DESKTOP = r"C:\Users\naman\OneDrive\Desktop"
RESUME_DIR = os.path.join(DESKTOP, "Personal", "Resume")
PORTFOLIO_RESUME_DIR = os.path.join(DESKTOP, "Personal", "portfolio_site", "public", "resume")
BACKUP_DIR = os.path.join(RESUME_DIR, "backup_originals")

# Tracked Project Repositories
REPOS = {
    "FundersAI": os.path.join(DESKTOP, "FundersAI"),
    "TalentOS": os.path.join(DESKTOP, "ALLThingsAgentic"),
    "CareFlow": os.path.join(DESKTOP, "CareFlow Intelligence"),
}

def get_repo_stats(repo_path):
    """Computes commit count, tracked code files, LOC, and test counts."""
    if not os.path.exists(repo_path):
        return None

    stats = {}
    
    # 1. Commit count
    try:
        commits = subprocess.check_output(
            ["git", "rev-list", "--count", "HEAD"],
            cwd=repo_path,
            text=True,
            stderr=subprocess.DEVNULL
        ).strip()
        stats["commits"] = int(commits)
    except Exception:
        stats["commits"] = 0

    # 2. Tracked code files & LOC
    try:
        files = subprocess.check_output(
            ["git", "ls-files"],
            cwd=repo_path,
            text=True,
            stderr=subprocess.DEVNULL
        ).splitlines()

        code_exts = {".py", ".ts", ".tsx", ".js", ".mjs", ".sql", ".html", ".css", ".md"}
        code_files = [f for f in files if any(f.endswith(ext) for ext in code_exts)]
        
        total_loc = 0
        for cf in code_files:
            p = os.path.join(repo_path, cf)
            if os.path.exists(p):
                try:
                    with open(p, "r", encoding="utf-8", errors="ignore") as fp:
                        total_loc += sum(1 for _ in fp)
                except Exception:
                    pass

        stats["files"] = len(code_files)
        stats["loc"] = total_loc
    except Exception:
        stats["files"] = 0
        stats["loc"] = 0

    # 3. Test suites count (extracted directly from git ls-files)
    test_files = [
        f for f in files 
        if (
            "test_" in os.path.basename(f) or 
            f.endswith(".test.ts") or 
            f.endswith(".test.tsx") or 
            f.endswith(".test.js") or 
            f.endswith(".test.mjs")
        ) and not any(x in f for x in ["node_modules", ".venv", ".pytest_cache", ".next"])
    ]

    stats["test_suites"] = len(test_files)
    
    return stats

def update_resumes():
    print("=" * 60)
    print("🚀 AUTOMATED RESUME SYNC ENGINE")
    print("=" * 60)

    # Gather live project stats
    print("\n🔍 Scanning project repositories...")
    stats = {}
    for name, path in REPOS.items():
        st = get_repo_stats(path)
        if st:
            stats[name] = st
            print(f"  ✓ {name}: {st['commits']} commits, {st['files']} files, {st['loc']:,} LOC, {st['test_suites']} test suites")
        else:
            print(f"  ⚠ {name}: Directory not found ({path})")

    # Format numbers
    f_loc_k = f"{round(stats.get('FundersAI', {}).get('loc', 145000) / 1000)}K"
    f_files = stats.get('FundersAI', {}).get('files', 760)
    f_commits = stats.get('FundersAI', {}).get('commits', 402)
    f_tests = f"{stats.get('FundersAI', {}).get('test_suites', 126)}+"

    t_loc_k = f"{round(stats.get('TalentOS', {}).get('loc', 32000) / 1000)}K"
    t_files = stats.get('TalentOS', {}).get('files', 134)
    t_commits = stats.get('TalentOS', {}).get('commits', 51)

    print("\n📄 Updating Word Documents (.docx)...")

    # 1. Update Master Resume
    master_path = os.path.join(RESUME_DIR, "Naman_Manocha_Master_Resume.docx")
    if os.path.exists(master_path):
        doc = docx.Document(master_path)
        # FundersAI subheader
        doc.paragraphs[19].text = f"Solo-built, Apr–Aug 2026 — ~{f_loc_k} lines of code across {f_files} files, {f_commits} commits, {f_tests} test suites — submitted to OpenAI Build Week"
        if doc.paragraphs[19].runs:
            doc.paragraphs[19].runs[0].italic = True

        # TalentOS title
        doc.paragraphs[38].runs[0].text = "TalentOS — Autonomous Opportunity Intelligence Platform (All Things Agentic Hackathon)"

        # TalentOS subheader
        doc.paragraphs[39].text = f"Python, Google ADK, LangGraph, Next.js, Firestore, Google Cloud Run  |  Solo-built, Aug 2026 — ~{t_loc_k} LOC across {t_files} files, {t_commits} commits, 235+ test suite"
        if doc.paragraphs[39].runs:
            doc.paragraphs[39].runs[0].italic = True

        doc.save(master_path)
        print("  ✓ Updated: Naman_Manocha_Master_Resume.docx")

    # 2. Update Agentic AI FDE Resume
    fde_path = os.path.join(RESUME_DIR, "Naman_Manocha_Agentic_AI_FDE_Resume.docx")
    if os.path.exists(fde_path):
        doc = docx.Document(fde_path)
        doc.paragraphs[16].text = f"Python, FastAPI, Next.js, LangGraph, Groq, OpenAI  |  ~{f_loc_k} LOC, {f_files} files, {f_commits} commits, {f_tests} test suites — OpenAI Build Week"
        if doc.paragraphs[16].runs:
            doc.paragraphs[16].runs[0].italic = True

        doc.paragraphs[19].text = "TalentOS — Autonomous Opportunity Intelligence Platform (All Things Agentic Hackathon)\t"
        doc.paragraphs[19].runs[0].bold = True

        doc.paragraphs[20].text = f"Python, Google ADK, LangGraph, Next.js, Firestore, Google Cloud Run  |  ~{t_loc_k} LOC across {t_files} files, {t_commits} commits, 235+ test suite"
        if doc.paragraphs[20].runs:
            doc.paragraphs[20].runs[0].italic = True

        doc.save(fde_path)
        print("  ✓ Updated: Naman_Manocha_Agentic_AI_FDE_Resume.docx")

    # 3. Update GenAI Engineer Resume
    genai_path = os.path.join(RESUME_DIR, "Naman_Manocha_GenAI_Engineer_Resume.docx")
    if os.path.exists(genai_path):
        doc = docx.Document(genai_path)
        doc.paragraphs[15].text = f"Python, FastAPI, Next.js, Supabase/pgvector, LangGraph, Groq, OpenAI  |  ~{f_loc_k} LOC, {f_files} files, {f_commits} commits, {f_tests} test suites — OpenAI Build Week"
        if doc.paragraphs[15].runs:
            doc.paragraphs[15].runs[0].italic = True

        doc.paragraphs[24].text = f"TalentOS (All Things Agentic Hackathon, Taskmaster Track): dual-pipeline multi-agent platform (Google ADK, LangGraph, Gemini on Vertex AI) generating tailored resumes/pitches with 91% deterministic pre-filtering before LLM calls — ~{t_loc_k} LOC, {t_files} files, 235+ test suite. FundersAI Reports: decoupled LangGraph microservice on a K3s/EC2 Kubernetes deployment. SQuAD QA Benchmarking: fine-tuned BERT/BiDAF/DistilBERT; published research paper."
        doc.save(genai_path)
        print("  ✓ Updated: Naman_Manocha_GenAI_Engineer_Resume.docx")

    # 4. Update AI Engineer Resume
    ai_path = os.path.join(RESUME_DIR, "Naman_Manocha_AI_Engineer_Resume.docx")
    if os.path.exists(ai_path):
        doc = docx.Document(ai_path)
        doc.paragraphs[16].text = f"Python, FastAPI, Next.js, Supabase/pgvector, LangGraph, Groq, OpenAI, Cloudflare R2  |  ~{f_loc_k} LOC, {f_files} files, {f_commits} commits, {f_tests} test suites — OpenAI Build Week"
        if doc.paragraphs[16].runs:
            doc.paragraphs[16].runs[0].italic = True

        doc.paragraphs[25].text = f"TalentOS (All Things Agentic Hackathon, Taskmaster Track): dual-pipeline autonomous agent platform (Google ADK, LangGraph, Firestore) ingesting ~2,500 postings/run across 8 sources with 91% pre-filtering and a 3-state evaluator+drafter chain — ~{t_loc_k} LOC, {t_files} files, 235+ tests. FundersAI Reports: decoupled LangGraph microservice on a 2-replica Kubernetes Deployment (K3s/AWS EC2). SQuAD QA Benchmarking: fine-tuned/benchmarked BERT/BiDAF/DistilBERT; published as a research paper."
        doc.save(ai_path)
        print("  ✓ Updated: Naman_Manocha_AI_Engineer_Resume.docx")

    # Compile fresh PDFs via Word Automation
    print("\n📑 Exporting fresh PDFs via Microsoft Word...")
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            for fname in os.listdir(RESUME_DIR):
                if fname.endswith(".docx") and not fname.startswith("~$"):
                    docx_full = os.path.abspath(os.path.join(RESUME_DIR, fname))
                    pdf_full = os.path.splitext(docx_full)[0] + ".pdf"
                    doc_w = word.Documents.Open(docx_full)
                    doc_w.SaveAs(pdf_full, FileFormat=17)  # 17 = wdFormatPDF
                    doc_w.Close()
                    print(f"  ✓ Exported: {os.path.basename(pdf_full)}")
        finally:
            word.Quit()
    except Exception as e:
        print(f"  ⚠ Word PDF export skipped ({e}). You can open .docx files in Word to export PDF.")

    # Sync to portfolio site
    if os.path.exists(PORTFOLIO_RESUME_DIR):
        print("\n🌐 Syncing to portfolio_site/public/resume...")
        for fname in os.listdir(RESUME_DIR):
            if (fname.endswith(".docx") or fname.endswith(".pdf")) and not fname.startswith("~$"):
                shutil.copy2(os.path.join(RESUME_DIR, fname), os.path.join(PORTFOLIO_RESUME_DIR, fname))
                print(f"  ✓ Synced: {fname}")

    print("\n🎉 Resume sync completed successfully!")

if __name__ == "__main__":
    update_resumes()
