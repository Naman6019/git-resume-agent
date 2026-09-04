import os
from typing import Dict, Any
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

def add_hyperlink(paragraph, url: str, text: str, color: str = "0066CC", underline: bool = True, italic: bool = True, bold: bool = False):
    """Appends an active, clickable hyperlink to a paragraph with styling."""
    if not url:
        return None
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    if italic:
        i = OxmlElement('w:i')
        rPr.append(i)

    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    new_run.append(rPr)
    
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

class DocxCompiler:
    """Modifies Word .docx resume templates with synthesized metrics and clickable hyperlinks."""

    def format_paragraph_with_links(self, paragraph, base_text: str, stats: Dict[str, Any], is_italic: bool = True):
        """Populates paragraph with base text and appends active clickable hyperlinks."""
        paragraph.text = base_text
        if paragraph.runs:
            paragraph.runs[0].italic = is_italic

        is_deployed = stats.get("deployed", False)
        live_url = stats.get("live_url")
        repo_url = stats.get("repo_url")

        if (is_deployed and live_url) or repo_url:
            sep_run = paragraph.add_run("  |  ")
            sep_run.italic = is_italic

            added = 0
            if is_deployed and live_url:
                lbl = paragraph.add_run("Live: ")
                lbl.italic = is_italic
                display_live = live_url.replace("https://", "").replace("http://", "").rstrip("/")
                if len(display_live) > 35:
                    display_live = display_live[:32] + "..."
                add_hyperlink(paragraph, live_url, display_live, color="0066CC", underline=True, italic=is_italic)
                added += 1

            if repo_url:
                if added > 0:
                    mid = paragraph.add_run(" | ")
                    mid.italic = is_italic
                lbl = paragraph.add_run("Repo: ")
                lbl.italic = is_italic
                display_repo = repo_url.replace("https://", "").replace("http://", "").rstrip("/")
                add_hyperlink(paragraph, repo_url, display_repo, color="0066CC", underline=True, italic=is_italic)

    def update_resume(self, file_path: str, persona_id: str, repo_stats: Dict[str, Dict[str, Any]]) -> bool:
        if not os.path.exists(file_path):
            return False

        doc = docx.Document(file_path)
        f_stats = repo_stats.get("FundersAI", {})
        t_stats = repo_stats.get("TalentOS", {})
        c_stats = repo_stats.get("CareFlow", {})
        g_stats = repo_stats.get("GitResume", {})

        f_loc = f_stats.get("loc_k", "145K")
        f_files = f_stats.get("files", 760)
        f_commits = f_stats.get("commits", 402)
        f_tests = f"{f_stats.get('test_suites', 120)}+"

        t_loc = t_stats.get("loc_k", "32K")
        t_files = t_stats.get("files", 134)
        t_commits = t_stats.get("commits", 51)

        if persona_id in ("master", "master_2page"):
            # P19: FundersAI subheader
            if len(doc.paragraphs) > 19:
                f_base = f"Solo-built, Apr–Aug 2026 — ~{f_loc} lines of code across {f_files} files, {f_commits} commits, {f_tests} test suites — submitted to OpenAI Build Week"
                self.format_paragraph_with_links(doc.paragraphs[19], f_base, f_stats, is_italic=True)

            # P38 & P39: TalentOS
            if len(doc.paragraphs) > 39:
                doc.paragraphs[38].runs[0].text = "TalentOS — Autonomous Opportunity Intelligence Platform (All Things Agentic Hackathon)"
                t_base = f"Python, Google ADK, LangGraph, Next.js, Firestore, Google Cloud Run  |  Solo-built, Aug 2026 — ~{t_loc} LOC across {t_files} files, {t_commits} commits, 235+ test suite"
                self.format_paragraph_with_links(doc.paragraphs[39], t_base, t_stats, is_italic=True)

        elif persona_id == "master_1page":
            # P16: FundersAI subheader
            if len(doc.paragraphs) > 16:
                f_base = f"Python, FastAPI, Next.js, Supabase/pgvector, LangGraph, Groq, OpenAI, Cloudflare R2, Kubernetes (K3s)\nSolo-built, Apr–Aug 2026 — ~{f_loc} LOC across {f_files} files, {f_commits} commits, {f_tests} test suites — OpenAI Build Week"
                self.format_paragraph_with_links(doc.paragraphs[16], f_base, f_stats, is_italic=True)

            # P21: TalentOS
            if len(doc.paragraphs) > 21:
                t_base = f"Python, Google ADK, LangGraph, Next.js, Firestore, Google Cloud Run  |  ~{t_loc} LOC across {t_files} files, {t_commits} commits, 235+ test suite"
                self.format_paragraph_with_links(doc.paragraphs[21], t_base, t_stats, is_italic=True)

        elif persona_id == "fde":
            if len(doc.paragraphs) > 16:
                f_base = f"Python, FastAPI, Next.js, LangGraph, Groq, OpenAI  |  ~{f_loc} LOC, {f_files} files, {f_commits} commits, {f_tests} test suites — OpenAI Build Week"
                self.format_paragraph_with_links(doc.paragraphs[16], f_base, f_stats, is_italic=True)
            if len(doc.paragraphs) > 20:
                doc.paragraphs[19].text = "TalentOS — Autonomous Opportunity Intelligence Platform (All Things Agentic Hackathon)\t"
                doc.paragraphs[19].runs[0].bold = True
                t_base = f"Python, Google ADK, LangGraph, Next.js, Firestore, Google Cloud Run  |  ~{t_loc} LOC across {t_files} files, {t_commits} commits, 235+ test suite"
                self.format_paragraph_with_links(doc.paragraphs[20], t_base, t_stats, is_italic=True)

        elif persona_id == "genai":
            if len(doc.paragraphs) > 15:
                f_base = f"Python, FastAPI, Next.js, Supabase/pgvector, LangGraph, Groq, OpenAI  |  ~{f_loc} LOC, {f_files} files, {f_commits} commits, {f_tests} test suites — OpenAI Build Week"
                self.format_paragraph_with_links(doc.paragraphs[15], f_base, f_stats, is_italic=True)
            if len(doc.paragraphs) > 24:
                t_links_str = " | " + t_stats.get("formatted_links", "") if t_stats.get("formatted_links") else ""
                doc.paragraphs[24].text = f"TalentOS (All Things Agentic Hackathon, Taskmaster Track): dual-pipeline multi-agent platform (Google ADK, LangGraph, Gemini on Vertex AI) generating tailored resumes/pitches with 91% deterministic pre-filtering before LLM calls — ~{t_loc} LOC, {t_files} files, 235+ test suite{t_links_str}. FundersAI Reports: decoupled LangGraph microservice on a K3s/EC2 Kubernetes deployment. SQuAD QA Benchmarking: fine-tuned BERT/BiDAF/DistilBERT; published research paper."

        elif persona_id == "ai_engineer":
            if len(doc.paragraphs) > 16:
                f_base = f"Python, FastAPI, Next.js, Supabase/pgvector, LangGraph, Groq, OpenAI, Cloudflare R2  |  ~{f_loc} LOC, {f_files} files, {f_commits} commits, {f_tests} test suites — OpenAI Build Week"
                self.format_paragraph_with_links(doc.paragraphs[16], f_base, f_stats, is_italic=True)
            if len(doc.paragraphs) > 25:
                t_links_str = " | " + t_stats.get("formatted_links", "") if t_stats.get("formatted_links") else ""
                doc.paragraphs[25].text = f"TalentOS (All Things Agentic Hackathon, Taskmaster Track): dual-pipeline autonomous agent platform (Google ADK, LangGraph, Firestore) ingesting ~2,500 postings/run across 8 sources with 91% pre-filtering and a 3-state evaluator+drafter chain — ~{t_loc} LOC, {t_files} files, 235+ tests{t_links_str}. FundersAI Reports: decoupled LangGraph microservice on a 2-replica Kubernetes Deployment (K3s/AWS EC2). SQuAD QA Benchmarking: fine-tuned/benchmarked BERT/BiDAF/DistilBERT; published as a research paper."

        doc.save(file_path)
        return True
